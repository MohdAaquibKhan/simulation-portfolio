"""
cv_builder_utils.py — Shared utilities for Aaquib's CV & Cover Letter generator scripts.
Claude will import from this module when generating tailored documents.
Place this file at D:\\Resume\\cv_builder_utils.py

Usage in generated scripts:
    import sys; sys.path.insert(0, r'D:\\Resume')
    from cv_builder_utils import *
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.section import WD_ORIENT
from docx.opc.constants import RELATIONSHIP_TYPE as RT
import datetime, os, io

# ── Palette ───────────────────────────────────────────────────────────────────
NAVY   = RGBColor(0x1F, 0x3A, 0x5F)
BLUE   = RGBColor(0x2E, 0x60, 0x9E)
GRAY   = RGBColor(0x55, 0x55, 0x55)
BLACK  = RGBColor(0x00, 0x00, 0x00)

# ── Personal constants ────────────────────────────────────────────────────────
FULL_NAME     = "Mohd Aaquib Khan"
EMAIL         = "a.khan.iitp@gmail.com"
PHONE         = "+91-9739107468"
LINKEDIN      = "linkedin.com/in/mohdaaquibkhan"
LOCATION      = "Bangalore, India"
DOB           = "23 June 1991, India"
NATIONALITY   = "Indian"
VISA          = "US B1 Visa (valid till Jun 2026)"   # Schengen expired Apr 2025 — omit
PHOTO_PATH    = r"D:\Resume\Aaquib_Photo.jpg"


def new_doc(margins_cm=2.0):
    """Create a new Document with standard margins."""
    doc = Document()
    for sec in doc.sections:
        m = Cm(margins_cm)
        sec.top_margin = sec.bottom_margin = sec.left_margin = sec.right_margin = m
    return doc


def set_spacing(para, before=0, after=0, line_pt=None, exact=False):
    pf = para.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after  = Pt(after)
    if line_pt is not None:
        if exact:
            pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        pf.line_spacing = Pt(line_pt)


def add_bottom_border(para, color="1F3A5F", sz=6):
    """Add a bottom border line to a paragraph (used for section dividers)."""
    pPr  = para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bot  = OxmlElement('w:bottom')
    bot.set(qn('w:val'),   'single')
    bot.set(qn('w:sz'),    str(sz))
    bot.set(qn('w:space'), '1')
    bot.set(qn('w:color'), color)
    pBdr.append(bot)
    pPr.append(pBdr)


def section_head(doc, text, before=10, after=2, color=NAVY, border_color="1F3A5F", font_size=10.5):
    """Bold section heading with bottom border — ATS safe (no table, no text box)."""
    p = doc.add_paragraph()
    set_spacing(p, before=before, after=after)
    add_bottom_border(p, color=border_color)
    r = p.add_run(text.upper())
    r.bold = True
    r.font.size = Pt(font_size)
    r.font.color.rgb = color
    return p


def job_header(doc, company, title, dates, location, before=8):
    """Two-line job block: company+dates on line 1, title+location italic on line 2."""
    p = doc.add_paragraph()
    set_spacing(p, before=before, after=0)
    p.paragraph_format.tab_stops.add_tab_stop(Inches(6.0), 2)  # right-align dates
    r1 = p.add_run(company); r1.bold = True; r1.font.size = Pt(10.5); r1.font.color.rgb = NAVY
    r2 = p.add_run(f"\t{dates}"); r2.font.size = Pt(10)
    p2 = doc.add_paragraph()
    set_spacing(p2, before=0, after=2)
    r3 = p2.add_run(f"{title}  •  {location}"); r3.italic = True; r3.font.size = Pt(10)
    return p


def add_bullet(doc, text, level=0, before=1, after=1):
    """ATS-safe bullet using docx numbering (not unicode character)."""
    p = doc.add_paragraph(style='List Bullet')
    set_spacing(p, before=before, after=after)
    p.paragraph_format.left_indent       = Pt(18 + level * 14)
    p.paragraph_format.first_line_indent = Pt(-9)
    r = p.add_run(text); r.font.size = Pt(10)
    return p


def add_plain(doc, text, size=10, bold=False, italic=False,
              before=2, after=2, indent=0, color=None, align=None):
    p = doc.add_paragraph()
    set_spacing(p, before=before, after=after)
    p.paragraph_format.left_indent = Pt(indent)
    if align:
        p.alignment = align
    r = p.add_run(text)
    r.font.size = Pt(size); r.bold = bold; r.italic = italic
    if color:
        r.font.color.rgb = color
    return p


def label_val(doc, label, value, size=10, before=1, after=1):
    """Bold label + normal value on same paragraph."""
    p = doc.add_paragraph()
    set_spacing(p, before=before, after=after)
    r1 = p.add_run(label); r1.bold = True; r1.font.size = Pt(size)
    r2 = p.add_run(value); r2.font.size = Pt(size)
    return p


def today_str(fmt="international"):
    """Return today's date formatted for CV/cover letter."""
    d = datetime.date.today()
    if fmt == "german":
        return d.strftime("%d.%m.%Y")
    else:
        return d.strftime("%d %B %Y")


def region_from_location(location_or_country):
    """
    Map a city or country string to a broad region folder name.

    Europe  : Germany, Austria, Switzerland, France, Netherlands, UK, Italy,
              Spain, Sweden, Finland, Norway, Denmark, Belgium, Poland, Czech,
              Romania, Hungary, Portugal, Greece, Ireland, Luxembourg,
              + common cities: Munich, Berlin, Stuttgart, Hamburg, Frankfurt,
                               Zurich, Vienna, Paris, London, Amsterdam, Milan,
                               Turin, Madrid, Barcelona, Stockholm, Helsinki,
                               Lausanne, Geneva, Giengen, Ulm, Schaan, Dortmund,
                               Cologne, Düsseldorf, Dresden, Leipzig, Nuremberg
    GCC     : UAE, Dubai, Abu Dhabi, Saudi Arabia, Riyadh, Qatar, Doha,
              Kuwait, Bahrain, Oman, Muscat, Sharjah
    India   : Bangalore, Mumbai, Delhi, Hyderabad, Chennai, Pune, Noida,
              Bengaluru, India
    USA     : USA, United States, New York, San Francisco, Seattle,
              Austin, Boston, Chicago, Conyers, Georgia
    Australia: Australia, Sydney, Melbourne, Brisbane
    """
    s = location_or_country.lower()

    europe_keys = [
        "germany","austria","switzerland","france","netherlands","uk","united kingdom",
        "italy","spain","sweden","finland","norway","denmark","belgium","poland",
        "czech","romania","hungary","portugal","greece","ireland","luxembourg",
        "munich","münchen","berlin","stuttgart","hamburg","frankfurt","zurich","zürich",
        "vienna","wien","paris","london","amsterdam","milan","milano","turin","torino",
        "madrid","barcelona","stockholm","helsinki","lausanne","geneva","genf",
        "giengen","ulm","schaan","dortmund","cologne","köln","düsseldorf","dresden",
        "leipzig","nuremberg","nürnberg","essen","bremen","hannover","bonn",
        "mannheim","karlsruhe","augsburg","wiesbaden","gelsenkirchen","mönchengladbach",
        "starnberg","ingolstadt","regensburg","wolfsburg","braunschweig","freiburg",
        "kiel","lübeck","erfurt","rostock","mainz","saarbrücken","potsdam",
        "europe","eu","emea"
    ]
    gcc_keys = [
        "uae","dubai","abu dhabi","abudhabi","saudi","riyadh","qatar","doha",
        "kuwait","bahrain","oman","muscat","sharjah","gcc","middle east"
    ]
    india_keys = [
        "india","bangalore","bengaluru","mumbai","delhi","hyderabad","chennai",
        "pune","noida","gurugram","gurgaon","kolkata","ahmedabad","jaipur"
    ]
    usa_keys = [
        "usa","united states","new york","san francisco","seattle","austin",
        "boston","chicago","conyers","georgia","california","texas","washington"
    ]
    australia_keys = ["australia","sydney","melbourne","brisbane","perth","canberra"]

    for k in gcc_keys:
        if k in s: return "GCC"
    for k in europe_keys:
        if k in s: return "Europe"
    for k in india_keys:
        if k in s: return "India"
    for k in usa_keys:
        if k in s: return "USA"
    for k in australia_keys:
        if k in s: return "Australia"
    return "Other"          # fallback — still creates a valid folder


def make_output_path(company_name, location, brief_role, doc_type="Resume", region=None):
    """
    Build and create the output folder, return full file path.

    Folder structure: D:\\Resume\\Applied\\[Region]\\[CompanyName]\\[Location_BriefRole]\\
    e.g.  D:\\Resume\\Applied\\Europe\\Siemens\\Munich_SrSimulationEngineer\\
          D:\\Resume\\Applied\\India\\Microsoft\\Bangalore_PrincipalEnggManager\\
          D:\\Resume\\Applied\\GCC\\Serco\\AbuDhabi_MechDesignerFEA\\

    - Region folder groups all applications by geography — easy to browse.
    - The company-level folder is REUSED if it already exists (exist_ok=True).
    - A new sub-folder is created for each position (location + role).

    Args:
        company_name : e.g. "Siemens Energy"
        location     : e.g. "Munich" or "Bangalore" (city only, no spaces)
        brief_role   : e.g. "SrSimulationEngineer" or "CAEManager" (no spaces)
        doc_type     : 'Resume' | 'Lebenslauf' | 'CoverLetter'
        region       : e.g. "Europe", "India", "GCC" — auto-detected if None

    File naming convention (name first, type second):
        Aaquib_Resume_[Company].docx
        Aaquib_CoverLetter_[Company].docx
        Aaquib_Lebenslauf_[Company].docx
    """
    def safe(s):
        return s.strip().replace(" ", "_").replace("/", "_").replace("\\", "_")

    co     = safe(company_name)
    sub    = f"{safe(location)}_{safe(brief_role)}"
    reg    = region if region else region_from_location(location)
    # exist_ok=True → region + company folders are reused; sub-folder is new per position
    folder = rf"D:\Resume\Applied\{reg}\{co}\{sub}"
    os.makedirs(folder, exist_ok=True)

    if doc_type == "Resume":
        fname = f"Aaquib_Resume_{co}.docx"
    elif doc_type == "Lebenslauf":
        fname = f"Aaquib_Lebenslauf_{co}.docx"
    else:                          # CoverLetter
        fname = f"Aaquib_CoverLetter_{co}.docx"
    return os.path.join(folder, fname)


def save_jd_docx(folder, jd_text, career_link=None, company_name="", job_title=""):
    """
    Write a JD.docx into the position folder, preserving the original job
    description text plus the career-portal link for future reference.

    Args:
        folder       : target folder (use os.path.dirname(out_cv))
        jd_text      : the full job-description text (pasted verbatim)
        career_link  : the job posting URL, if provided (None to skip)
        company_name : e.g. "ExxonMobil" — for the heading
        job_title    : e.g. "Computational Scientist" — for the heading
    """
    os.makedirs(folder, exist_ok=True)
    doc = Document()
    for sec in doc.sections:
        m = Cm(2.0)
        sec.top_margin = sec.bottom_margin = sec.left_margin = sec.right_margin = m

    head = doc.add_paragraph()
    r = head.add_run("Job Description")
    r.bold = True
    r.font.size = Pt(16)
    r.font.color.rgb = NAVY

    if company_name or job_title:
        sub = doc.add_paragraph()
        sr = sub.add_run(" — ".join([x for x in (job_title, company_name) if x]))
        sr.bold = True
        sr.font.size = Pt(11)
        sr.font.color.rgb = GRAY

    if career_link:
        lp = doc.add_paragraph()
        lp.add_run("Career link: ").bold = True
        add_hyperlink(lp, career_link, career_link, size=10)

    doc.add_paragraph()  # spacer
    for line in jd_text.split("\n"):
        p = doc.add_paragraph()
        set_spacing(p, before=0, after=2)
        p.add_run(line).font.size = Pt(10)

    path = os.path.join(folder, "JD.docx")
    doc.save(path)
    return path


def save_as_pdf(docx_path):
    """
    Convert a .docx file to PDF in the same folder.
    Uses docx2pdf (wraps Microsoft Word on Windows via COM, LibreOffice elsewhere).
    Returns the PDF path, or None if conversion fails.
    """
    try:
        from docx2pdf import convert
        pdf_path = docx_path.replace(".docx", ".pdf")
        convert(docx_path, pdf_path)
        return pdf_path
    except Exception as e:
        print(f"  [PDF conversion skipped: {e}]")
        return None


def add_subheading(doc, text, before=5, after=1):
    """Bold italic sub-section label within an experience block."""
    p = doc.add_paragraph()
    set_spacing(p, before=before, after=after)
    r = p.add_run(text)
    r.bold = True; r.italic = True; r.font.size = Pt(10); r.font.color.rgb = BLUE
    return p


def name_header_intl(doc, title_line, include_photo=False, photo_path=None):
    """
    International CV: large name + subtitle + contact line with clickable LinkedIn.
    include_photo=True → adds a passport photo top-right (recommended for Germany even in English format).
    """
    if include_photo:
        # Two-column layout: text left (~13cm), photo right (~3.5cm)
        tbl = doc.add_table(rows=1, cols=2)
        tbl.style = 'Table Grid'
        tbl_pr = OxmlElement('w:tblPr')
        tbl_bdr = OxmlElement('w:tblBorders')
        for side in ('top','left','bottom','right','insideH','insideV'):
            el = OxmlElement(f'w:{side}')
            el.set(qn('w:val'), 'none')
            tbl_bdr.append(el)
        tbl_pr.append(tbl_bdr)
        tbl._tbl.insert(0, tbl_pr)
        tbl.rows[0].cells[0].width = Cm(13)
        tbl.rows[0].cells[1].width = Cm(3.8)
        text_cell  = tbl.rows[0].cells[0]
        photo_cell = tbl.rows[0].cells[1]
        text_cell._tc.clear_content()

        def _p(cell):
            p = cell.add_paragraph(); return p

        pn = _p(text_cell); set_spacing(pn, before=0, after=2)
        r  = pn.add_run(FULL_NAME)
        r.bold = True; r.font.size = Pt(18); r.font.color.rgb = NAVY

        pt = _p(text_cell); set_spacing(pt, before=0, after=1)
        pt.add_run(title_line).font.size = Pt(10)
        pt.runs[0].font.color.rgb = GRAY

        pc = _p(text_cell); set_spacing(pc, before=0, after=4)
        pc.add_run(f"{EMAIL}  |  {PHONE}  |  ").font.size = Pt(9.5)
        add_hyperlink(pc, LINKEDIN, f"https://{LINKEDIN}", size=9.5)
        pc.add_run(f"  |  {LOCATION}").font.size = Pt(9.5)

        insert_photo_in_cell(photo_cell, photo_path or PHOTO_PATH, width_cm=3.3)
    else:
        p = doc.add_paragraph()
        set_spacing(p, before=0, after=2)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(FULL_NAME)
        r.bold = True; r.font.size = Pt(18); r.font.color.rgb = NAVY

        p2 = doc.add_paragraph()
        set_spacing(p2, before=0, after=1)
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r2 = p2.add_run(title_line)
        r2.font.size = Pt(10); r2.font.color.rgb = GRAY

        p3 = doc.add_paragraph()
        set_spacing(p3, before=0, after=4)
        p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p3.add_run(f"{EMAIL}  |  {PHONE}  |  ").font.size = Pt(9.5)
        add_hyperlink(p3, LINKEDIN, f"https://{LINKEDIN}", size=9.5)
        p3.add_run(f"  |  {LOCATION}").font.size = Pt(9.5)


def add_education_entry(doc, degree, institution, dates, cgpa=None, details=None):
    p = doc.add_paragraph()
    set_spacing(p, before=6, after=0)
    p.paragraph_format.tab_stops.add_tab_stop(Inches(6.0), 2)
    r1 = p.add_run(degree); r1.bold = True; r1.font.size = Pt(10)
    p.add_run(f"\t{dates}").font.size = Pt(10)
    p2 = doc.add_paragraph()
    set_spacing(p2, before=0, after=1)
    inst_str = institution + (f"  |  CGPA: {cgpa}" if cgpa else "")
    r3 = p2.add_run(inst_str); r3.italic = True; r3.font.size = Pt(10)
    if details:
        for d in details:
            add_bullet(doc, d)


# ── Cover Letter Helpers ──────────────────────────────────────────────────────
def cover_letter_header(doc, company_name, job_title, date_fmt="international",
                        manager_name=None, company_address=None, ref_number=None):
    """Build the top block of a cover letter."""
    # Sender block
    add_plain(doc, FULL_NAME, size=12, bold=True, before=0, after=1)
    add_plain(doc, f"{LOCATION}  |  {PHONE}  |  {EMAIL}", size=10, before=0, after=1)
    p_li = doc.add_paragraph()
    set_spacing(p_li, before=0, after=8)
    add_hyperlink(p_li, LINKEDIN, f"https://{LINKEDIN}", size=10)

    # Date
    add_plain(doc, today_str(date_fmt), size=10, before=0, after=6)

    # Recipient
    add_plain(doc, company_name, size=10, bold=True, before=0, after=0)
    if company_address:
        add_plain(doc, company_address, size=10, before=0, after=6)
    else:
        doc.add_paragraph()  # blank line

    # Subject
    subj = f"Application for {job_title}"
    if ref_number:
        subj += f" — Ref: {ref_number}"
    p = doc.add_paragraph()
    set_spacing(p, before=0, after=6)
    r = p.add_run(f"Subject: {subj}")
    r.bold = True; r.font.size = Pt(10)

    # Salutation
    salutation = f"Dear {manager_name}," if manager_name else "Dear Hiring Team,"
    add_plain(doc, salutation, size=10, before=0, after=4)


def cover_letter_paragraph(doc, text, before=0, after=6):
    p = doc.add_paragraph()
    set_spacing(p, before=before, after=after, line_pt=13)
    r = p.add_run(text); r.font.size = Pt(10.5)
    return p


def cover_letter_closing(doc, date_fmt="international", lang="en"):
    if lang == "de":
        sign_off = "Mit freundlichen Grüßen,"
    else:
        sign_off = "Kind regards,"
    add_plain(doc, sign_off, size=10.5, before=6, after=20)
    add_plain(doc, FULL_NAME, size=10.5, bold=True, before=0, after=2)
    add_plain(doc, "Enclosure: Curriculum Vitae", size=10, italic=True, before=0, after=0)


# ── Hyperlink helper ─────────────────────────────────────────────────────────
def add_hyperlink(para, text, url, size=10, color=None, underline=True):
    """
    Add a clickable hyperlink run to an existing paragraph.
    Works in any paragraph — use for LinkedIn, email, websites.

    Usage:
        p = doc.add_paragraph()
        add_hyperlink(p, "linkedin.com/in/mohdaaquibkhan",
                      "https://linkedin.com/in/mohdaaquibkhan")
    """
    part  = para.part
    r_id  = part.relate_to(
        url,
        'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink',
        is_external=True
    )
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)

    new_run = OxmlElement('w:r')
    rPr     = OxmlElement('w:rPr')

    # Style
    rStyle = OxmlElement('w:rStyle')
    rStyle.set(qn('w:val'), 'Hyperlink')
    rPr.append(rStyle)

    # Override size / colour if requested
    sz_el = OxmlElement('w:sz')
    sz_el.set(qn('w:val'), str(int(size * 2)))
    rPr.append(sz_el)

    if color:
        col_el = OxmlElement('w:color')
        col_el.set(qn('w:val'), f"{color.rgb.red:02X}{color.rgb.green:02X}{color.rgb.blue:02X}")
        rPr.append(col_el)

    if not underline:
        u_el = OxmlElement('w:u')
        u_el.set(qn('w:val'), 'none')
        rPr.append(u_el)

    new_run.append(rPr)
    t_el = OxmlElement('w:t')
    t_el.text = text
    new_run.append(t_el)
    hyperlink.append(new_run)
    para._p.append(hyperlink)
    return hyperlink


# ── Photo helper ──────────────────────────────────────────────────────────────
def _make_placeholder_photo():
    """
    Create an in-memory placeholder image (light grey with 'Foto' text).
    Used when the real photo file is not yet available.
    """
    from PIL import Image, ImageDraw, ImageFont
    W, H = 280, 360          # ~3.5 × 4.5 cm at 80 dpi
    img  = Image.new("RGB", (W, H), color=(220, 220, 220))
    draw = ImageDraw.Draw(img)
    draw.rectangle([0, 0, W-1, H-1], outline=(160, 160, 160), width=2)
    draw.text((W//2, H//2), "Foto", fill=(100, 100, 100), anchor="mm")
    buf = io.BytesIO()
    img.save(buf, format="JPEG")
    buf.seek(0)
    return buf


def insert_photo_in_cell(cell, photo_path=None, width_cm=3.5):
    """
    Insert a professional photo into a table cell, right-aligned.
    If photo_path is None or does not exist, inserts a placeholder box.

    Standard German CV photo size: ~3.5 cm wide × 4.5 cm tall (passport-style).

    Usage (German CV header table):
        insert_photo_in_cell(tbl.rows[0].cells[1], PHOTO_PATH)
    """
    para = cell.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_spacing(para, before=0, after=0)

    run = para.add_run()
    if photo_path and os.path.exists(photo_path):
        run.add_picture(photo_path, width=Cm(width_cm))
    else:
        buf = _make_placeholder_photo()
        run.add_picture(buf, width=Cm(width_cm))
        # Add a note below so the user knows what to do
        note = cell.add_paragraph()
        set_spacing(note, before=2, after=0)
        note.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = note.add_run("← save photo as\nD:\\Resume\\Aaquib_Photo.jpg")
        r.font.size = Pt(7); r.font.color.rgb = RGBColor(0x99, 0x00, 0x00)


# ── German CV header block ────────────────────────────────────────────────────
def lebenslauf_header(doc, photo_path=None):
    """
    Build the standard German Lebenslauf header:
      - 'LEBENSLAUF' centred title
      - 2-column borderless table: personal details LEFT, photo RIGHT
    Photo column is ~4 cm wide; details column takes the rest.
    """
    # Title
    p_title = doc.add_paragraph()
    set_spacing(p_title, before=0, after=8)
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p_title.add_run("LEBENSLAUF")
    r.bold = True; r.font.size = Pt(18); r.font.color.rgb = NAVY

    # Header table — 2 cols, no borders
    tbl = doc.add_table(rows=1, cols=2)
    tbl.style = 'Table Grid'

    # Remove all borders from the table
    tbl_pr = OxmlElement('w:tblPr')
    tbl_bdr = OxmlElement('w:tblBorders')
    for side in ('top','left','bottom','right','insideH','insideV'):
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'), 'none')
        tbl_bdr.append(el)
    tbl_pr.append(tbl_bdr)
    tbl._tbl.insert(0, tbl_pr)

    # Column widths: details ~11cm, photo ~4cm  (A4 - 4cm margins = ~17cm content)
    details_w = Cm(11)
    photo_w   = Cm(4)
    from docx.oxml import OxmlElement as OE
    tbl.rows[0].cells[0].width = details_w
    tbl.rows[0].cells[1].width = photo_w

    cell_d = tbl.rows[0].cells[0]
    cell_p = tbl.rows[0].cells[1]

    # Vertical align top for both cells
    for c in (cell_d, cell_p):
        tc_pr = c._tc.get_or_add_tcPr()
        va = OxmlElement('w:vAlign'); va.set(qn('w:val'), 'top')
        tc_pr.append(va)

    # ── Personal details (left cell) ──────────────────────────────────────────
    def det(label, value, url=None):
        """One details row: bold label + value (optionally hyperlinked)."""
        p = cell_d.add_paragraph()
        set_spacing(p, before=1, after=1)
        r1 = p.add_run(f"{label}: "); r1.bold = True; r1.font.size = Pt(10)
        if url:
            add_hyperlink(p, value, url, size=10)
        else:
            r2 = p.add_run(value); r2.font.size = Pt(10)

    # Remove the default empty paragraph in cell_d
    cell_d._tc.clear_content()

    det("Name",                FULL_NAME)
    det("Geburtsdatum",        "23. Juni 1991, Indien")
    det("Staatsangehörigkeit", "Indisch")
    det("Adresse",             "Bangalore, Indien")
    det("Telefon",             PHONE)
    det("E-Mail",              EMAIL,    url=f"mailto:{EMAIL}")
    det("LinkedIn",            LINKEDIN, url=f"https://{LINKEDIN}")
    det("US B1-Visum",         "Gültig bis Jun 2026")   # Schengen abgelaufen — weggelassen

    # ── Photo (right cell) ────────────────────────────────────────────────────
    insert_photo_in_cell(cell_p, photo_path, width_cm=3.5)

    return tbl


print("cv_builder_utils loaded successfully.")
